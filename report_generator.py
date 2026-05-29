"""
v3 보고서 생성기
- 분석 인사이트가 각 섹션에 인라인 배치
- 매체별 작품 구성 포함
- 평가 초안 자동 반영
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import tempfile

from styles import (
    setup_document, set_run_font, add_paragraph, add_horizontal_rule,
    add_section_title, add_subsection_title, add_sub2_title,
    add_detail_title, add_bullet_main, add_bullet_sub, add_arrow_note,
    create_table, create_table_left_aligned,
    add_image, add_images_auto, add_images_2col,
    add_page_break, add_page_numbers_right,
    Colors, Fonts, CIRCLED_NUMBERS, ImageSize,
)
from chart_generator import (
    create_weekly_visitors_chart,
    create_media_composition_chart,
)


class ExhibitionReportGenerator:

    def __init__(self, data):
        self.data = data
        self.doc = Document()
        self.temp_files = []
        # LLM이 재작성한 분석 문단 (없으면 룰 기반 폴백)
        self.llm_sections = data.get("llm_sections", {})

    def generate(self, output_path):
        setup_document(self.doc)
        add_page_numbers_right(self.doc)

        self._create_toc_page()
        add_page_break(self.doc)

        self._section_1_overview()
        self._section_2_theme()
        add_page_break(self.doc)

        self._section_3_composition()
        add_page_break(self.doc)

        self._section_4_results()

        if self._has_promotion_data():
            add_page_break(self.doc)
            self._section_5_promotion()

        add_page_break(self.doc)
        self._section_6_evaluation()

        add_paragraph(self.doc, "")
        add_paragraph(self.doc, "끝.", size=Fonts.BODY, bold=False,
                      alignment=WD_ALIGN_PARAGRAPH.LEFT,
                      space_before=Pt(12), space_after=Pt(0), line_spacing=1.15)

        self.doc.save(output_path)
        self._cleanup()
        return output_path

    def _cleanup(self):
        for f in self.temp_files:
            try: os.remove(f)
            except OSError: pass

    # ─── 인라인 분석 삽입 헬퍼 ───

    def _insert_summary_metrics_table(self):
        """VI. Executive Summary 상단 핵심 수치 종합표.

        디렉터가 이 섹션만 읽어도 핵심 사실 파악이 가능하도록
        본문 II~V장의 주요 정량 지표를 4열 표로 압축.
        평가어 없이 사실만 제시.
        """
        metrics = self.data.get("summary_metrics", [])
        if not metrics:
            return

        reference_label = metrics[0].get("reference_label", "역대 전시")

        add_subsection_title(self.doc, "1", "핵심 수치 종합")
        add_paragraph(
            self.doc,
            f"(비교 기준: {reference_label} 평균)",
            size=Fonts.CAPTION,
            color=Colors.MEDIUM_GRAY,
            space_after=Pt(4),
        )

        headers = ["지표", "본 전시", "비교 평균", "차이"]
        rows = [
            [m["label"], m["current_fmt"], m["reference_avg_fmt"], m["diff_fmt"]]
            for m in metrics
        ]
        create_table_left_aligned(
            self.doc,
            rows=len(rows),
            cols=4,
            data=rows,
            headers=headers,
            first_col_bold=True,
        )
        add_paragraph(self.doc, "", space_after=Pt(6))

    def _insert_section_insights(self, section_key):
        """해당 섹션에 분석 문단 삽입 — LLM 결과 우선, 없으면 룰 기반 폴백"""

        # 1) LLM이 재작성한 문단이 있으면 그것을 사용
        llm_text = self.llm_sections.get(section_key, "")
        if llm_text and llm_text.strip():
            add_paragraph(self.doc, "", space_before=Pt(6))
            for para_text in llm_text.split("\n\n"):
                para_text = para_text.strip()
                if para_text:
                    add_paragraph(self.doc, para_text, size=Fonts.BODY,
                                  space_after=Pt(6), line_spacing=1.5,
                                  first_line_indent=Cm(0.5))
            return

        # 2) 폴백: 룰 기반 인사이트를 화살표 노트로 삽입
        insights = self.data.get("section_insights", {}).get(section_key, [])
        if not insights:
            return

        add_paragraph(self.doc, "", space_before=Pt(4))

        grouped = {}
        for ins in insights:
            label = ins.get("category_label", ins.get("category", "분석"))
            grouped.setdefault(label, []).append(ins)

        for label, items in grouped.items():
            add_bullet_main(self.doc, None, f"[데이터 분석] {label}", bold_value=True)
            for item in items:
                add_arrow_note(self.doc, item["text"])

    # ─── 전시 구성 서술 문단 생성 ───

    def _build_composition_narrative(self):
        """실제 보고서의 '전시 구성 > 1. 전시' 서술 개요 문단을 자동 생성.

        예시 출력:
        "전시는 일민미술관 1, 2, 3전시실 및 프로젝트 룸에서 진행되었다.
         출품 작품은 회화 14점, 설치 6점, 조각 2점, 영상 2점 총 24점이다.
         그래픽 디자인은 페이퍼프레스가, 공간 구성은 석운동이 맡았다."
        """
        rooms = self.data.get("rooms", [])
        if not rooms:
            return ""

        # 1) 전시실 나열
        room_names = [r.get("name", "") for r in rooms if r.get("name")]
        if not room_names:
            return ""

        if len(room_names) == 1:
            venue_str = f"일민미술관 {room_names[0]}"
        else:
            venue_str = "일민미술관 " + ", ".join(room_names[:-1]) + " 및 " + room_names[-1]

        sentences = [f"전시는 {venue_str}에서 진행되었다."]

        # 2) 작품 매체 구성 (전체 합산)
        artworks = self.data.get("artworks", {})
        total = artworks.get("total", 0)
        if total and total > 0:
            media_order = [
                ("painting", "회화"), ("sculpture", "조각"), ("photo", "사진"),
                ("installation", "설치"), ("media", "영상"), ("other", "기타"),
            ]
            media_parts = []
            for key, label in media_order:
                v = artworks.get(key, 0)
                if v and v > 0:
                    media_parts.append(f"{label} {v}점")

            if media_parts:
                sentences.append(
                    f"출품 작품은 {', '.join(media_parts)} 총 {total}점이다."
                )

        # 3) 디자인 크레딧
        graphic = self.data.get("graphic_designer", "")
        space = self.data.get("space_designer", "")

        def _subj_particle(name):
            """이름 뒤에 '이/가' 조사를 적절히 붙임"""
            if not name:
                return name
            last_char = name[-1]
            code = ord(last_char)
            if 0xAC00 <= code <= 0xD7A3:
                return name + ("이" if (code - 0xAC00) % 28 != 0 else "가")
            # 한글이 아닌 경우 (영문 등) — 그냥 '가' 사용
            return name + "가"

        if graphic and space:
            sentences.append(
                f"그래픽 디자인은 {_subj_particle(graphic)}, "
                f"공간 구성은 {_subj_particle(space)} 맡았다."
            )
        elif graphic:
            sentences.append(f"그래픽 디자인은 {_subj_particle(graphic)} 맡았다.")
        elif space:
            sentences.append(f"공간 구성은 {_subj_particle(space)} 맡았다.")

        return " ".join(sentences)

    # ─── 목차 ───

    def _create_toc_page(self):
        title = self.data.get("exhibition_title", "전시 제목")
        add_paragraph(self.doc, f"전시보고서 - 《{title}》",
                      size=Fonts.TOC_TITLE, bold=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER,
                      space_before=Pt(12), space_after=Pt(4))
        add_horizontal_rule(self.doc)

        toc_items = [
            "I. 전시 개요",
            "II. 전시 주제와 내용",
            "III. 전시 구성",
            "IV. 전시 결과",
            "V. 홍보 방식 및 언론 보도",
            "VI. Executive Summary",
        ]
        for item in toc_items:
            add_paragraph(self.doc, item, size=Fonts.TOC_ITEM, bold=True,
                          space_before=Pt(3), space_after=Pt(3), line_spacing=1.15)
            add_horizontal_rule(self.doc)

        poster = self.data.get("poster_image")
        if poster and os.path.exists(poster):
            add_paragraph(self.doc, "", space_before=Pt(10))
            add_image(self.doc, poster, width=ImageSize.POSTER_WIDTH)

    # ─── I. 전시 개요 ───

    def _section_1_overview(self):
        add_section_title(self.doc, "I", "전시 개요")
        ov = self.data.get("overview", {})

        fields = [
            ("title", "전시 제목", lambda v: f"《{v}》"),
            ("period", "전시 기간", None),
            ("exhibition_days", "전시 일수", None),
            ("artists", "참여 작가", lambda v: ", ".join(v) if isinstance(v, list) else v),
            ("chief_curator", "책임기획", None),
            ("curators", "기획", None),
            ("coordinators", "진행", None),
            ("curatorial_team", "학예팀", None),
            ("pr", "홍보", None),
            ("sponsors", "후원", None),
        ]
        for key, label, fmt in fields:
            val = ov.get(key)
            if val:
                display = fmt(val) if fmt else val
                add_bullet_main(self.doc, label, display)

        # 예산 (굵은 + 밑줄)
        if ov.get("total_budget"):
            add_bullet_main(self.doc, "총 사용 예산", ov["total_budget"],
                            bold_value=True, underline_value=True)
            for item in ov.get("budget_breakdown", []):
                add_bullet_sub(self.doc, item)

        if ov.get("total_revenue"):
            add_bullet_main(self.doc, "총수입", ov["total_revenue"])
        if ov.get("programs"):
            add_bullet_main(self.doc, "프로그램", ov["programs"])
        if ov.get("staff_count"):
            add_bullet_main(self.doc, "운영 인력", ov["staff_count"])
        if ov.get("visitors"):
            add_bullet_main(self.doc, "관객 수", ov["visitors"],
                            bold_value=True, underline_value=True)

        # 출품 작품 (v3 신규)
        artworks = self.data.get("artworks", {})
        if artworks.get("total"):
            media_parts = []
            for key, label in [("painting", "회화"), ("sculpture", "조각"), ("photo", "사진"),
                                ("installation", "설치"), ("media", "미디어"), ("other", "기타")]:
                v = artworks.get(key, 0)
                if v: media_parts.append(f"{label} {v}점")
            overview_text = f"{artworks['total']}점"
            if media_parts:
                overview_text += f" ({', '.join(media_parts)})"
            add_bullet_main(self.doc, "출품 작품", overview_text)

        add_paragraph(self.doc, "")

    # ─── II. 전시 주제와 내용 ───

    def _section_2_theme(self):
        add_section_title(self.doc, "II", "전시 주제와 내용")
        theme = self.data.get("theme_text", "")
        if theme:
            for p_text in theme.split("\n\n"):
                p_text = p_text.strip()
                if p_text:
                    add_paragraph(self.doc, p_text, size=Fonts.BODY,
                                  space_after=Pt(6), line_spacing=1.5,
                                  first_line_indent=Cm(0.5))

    # ─── III. 전시 구성 ───

    def _section_3_composition(self):
        add_section_title(self.doc, "III", "전시 구성")

        self._sub_rooms()
        self._sub_artwork_composition()  # 매체·신작 도넛(자기완결형) + 서술
        self._sub_programs()
        self._sub_staff()
        self._sub_materials()

        # v3: 인라인 분석 (프로그램, 작품, 인력)
        self._insert_section_insights("composition")

    def _sub_artwork_composition(self):
        """출품 작품 구성 — 매체 도넛 + 신작/구작 도넛(2-up, 자기완결형).

        도넛은 한 줄에 2개. 이어서 그 도표에 대한 서술 1문단(원칙: 차트→서술).
        """
        art = self.data.get("artworks", {})
        total = art.get("total", 0) or 0
        if total <= 0:
            return

        media = {
            "회화": art.get("painting", 0), "조각": art.get("sculpture", 0),
            "사진": art.get("photo", 0), "설치": art.get("installation", 0),
            "미디어": art.get("media", 0), "기타": art.get("other", 0),
        }
        media_chart = create_media_composition_chart(media, title="출품 매체 구성")

        new_n = art.get("new", 0) or 0
        old_n = art.get("old", 0) or 0
        new_chart = None
        if new_n > 0 or old_n > 0:
            new_chart = create_media_composition_chart(
                {"신작": new_n, "구작": old_n}, title="신작·구작 구성")

        charts = [c for c in (media_chart, new_chart) if c]
        if not charts:
            return
        add_subsection_title(self.doc, "2", "출품 작품 구성")
        for c in charts:
            self.temp_files.append(c)
        if len(charts) == 2:
            add_images_2col(self.doc, charts)  # 한 줄에 2개
        else:
            add_image(self.doc, charts[0], is_chart=True)

        # 서술 (차트→서술 원칙)
        media_items = [(k, v) for k, v in media.items() if v]
        media_items.sort(key=lambda kv: -kv[1])
        parts = []
        if media_items:
            top = media_items[0]
            parts.append(
                f"출품작 {total}점의 매체 구성은 {top[0]}({top[1]}점, "
                f"{top[1]/total*100:.0f}%)의 비중이 가장 큼.")
        if new_n or old_n:
            parts.append(
                f"이 중 신작은 {new_n}점({new_n/total*100:.0f}%), "
                f"구작은 {old_n}점으로 구성됨.")
        if parts:
            add_paragraph(self.doc, " ".join(parts), size=Fonts.BODY,
                          space_after=Pt(6), line_spacing=1.5,
                          first_line_indent=Cm(0.5))

    def _sub_rooms(self):
        add_subsection_title(self.doc, "1", "전시")

        # ── 전시 구성 서술 개요 문단 자동 생성 ──
        overview_para = self._build_composition_narrative()
        if overview_para:
            add_paragraph(self.doc, overview_para, size=Fonts.BODY,
                          space_after=Pt(6), line_spacing=1.5,
                          first_line_indent=Cm(0.5))

        rooms = self.data.get("rooms", [])
        for i, room in enumerate(rooms):
            add_sub2_title(self.doc, i + 1, room.get("name", f"{i+1}전시실"))
            artists = room.get("artists", "")
            if artists:
                if isinstance(artists, list): artists = ", ".join(artists)
                add_detail_title(self.doc, CIRCLED_NUMBERS[0], "참여 작가")
                add_paragraph(self.doc, artists, size=Fonts.BODY, left_indent=Cm(0.8))
            fp = room.get("floor_plan")
            if fp and os.path.exists(fp):
                add_detail_title(self.doc, CIRCLED_NUMBERS[1], "도면")
                add_image(self.doc, fp)
            photos = room.get("photos", [])
            valid = [p for p in photos if os.path.exists(p)]
            if valid:
                idx = 2 if (fp and os.path.exists(fp)) else 1
                add_detail_title(self.doc, CIRCLED_NUMBERS[idx], "전경 사진")
                add_images_auto(self.doc, valid)

    def _sub_programs(self):
        programs = self.data.get("related_programs", [])
        total_count = len(programs)
        total_part = sum(int(str(p.get("participants", "0")).replace(",", "").replace("명", ""))
                         for p in programs if p.get("participants"))

        suffix = ""
        if total_count > 0:
            suffix = f" - 총 {total_count}개 프로그램 진행"
            if total_part > 0:
                suffix += f", {total_part:,}명 참여"

        add_subsection_title(self.doc, "2", "전시 연계 프로그램", suffix=suffix)

        if programs:
            headers = ["구분", "제목", "일자", "참여 인원", "비고"]
            table_data = [[p.get("category", ""), p.get("title", ""), p.get("date", ""),
                           p.get("participants", ""), p.get("note", "")] for p in programs]
            create_table(self.doc, len(table_data), 5, data=table_data, headers=headers,
                         col_widths=[Cm(2), Cm(5.5), Cm(2.5), Cm(1.5), Cm(4)])

    def _sub_staff(self):
        add_subsection_title(self.doc, "3", "전시 운영 인력")
        staff = self.data.get("staff", {})
        if staff.get("main_staff"):
            add_sub2_title(self.doc, "1", "스태프")
            info = staff["main_staff"]
            if isinstance(info, dict):
                if info.get("count"):
                    add_detail_title(self.doc, CIRCLED_NUMBERS[0], "인원")
                    add_paragraph(self.doc, info["count"], left_indent=Cm(0.8))
                if info.get("role"):
                    add_detail_title(self.doc, CIRCLED_NUMBERS[1], "역할 및 활동 내용")
                    add_paragraph(self.doc, info["role"], left_indent=Cm(0.8))

    def _sub_materials(self):
        add_subsection_title(self.doc, "4", "인쇄물 및 굿즈")
        materials = self.data.get("printed_materials", [])
        if materials:
            headers = ["종류", "제작 수량", "비고"]
            table_data = [[m.get("type", ""), m.get("quantity", ""), m.get("note", "")] for m in materials]
            create_table(self.doc, len(table_data), 3, data=table_data, headers=headers,
                         col_widths=[Cm(5.5), Cm(3), Cm(6.5)])

    # ─── IV. 전시 결과 ───

    def _section_4_results(self):
        add_section_title(self.doc, "IV", "전시 결과")
        self._sub_budget()
        self._sub_revenue()
        self._sub_visitor_composition()

        # v3: 인라인 분석 (예산, 관객)
        self._insert_section_insights("results")

    def _sub_budget(self):
        add_subsection_title(self.doc, "1", "예산 및 지출")
        budget = self.data.get("budget", {})
        if budget.get("total_spent"):
            add_bullet_main(self.doc, "지출 총액", budget["total_spent"],
                            bold_value=True, underline_value=True)
        for note in budget.get("breakdown_notes", []):
            add_bullet_sub(self.doc, note)

        summary = budget.get("summary", [])
        if summary:
            add_paragraph(self.doc, "", space_before=Pt(6))
            headers = ["사업", "계획 예산(원)", "집행 예산(원)", "계획 대비 집행"]
            table_data = [[s.get("category", ""), s.get("planned", ""),
                           s.get("actual", ""), s.get("note", "")] for s in summary]
            create_table(self.doc, len(table_data), 4, data=table_data, headers=headers,
                         col_widths=[Cm(2.5), Cm(4.5), Cm(4.5), Cm(4)])

        for note in budget.get("arrow_notes", []):
            add_arrow_note(self.doc, note)

        # 세부 내역 표 (엑셀 업로드로 입력된 상세 예산)
        details = budget.get("details", [])
        if details:
            add_paragraph(self.doc, "", space_before=Pt(8))
            add_bullet_main(self.doc, None, "예산 세부 내역", bold_value=True)
            add_paragraph(self.doc, "", space_before=Pt(2))
            headers = ["사업 구분", "항목", "세부 내용", "금액(원)", "비고"]
            table_data = [[d.get("category", ""), d.get("subcategory", ""),
                           d.get("detail", ""), d.get("amount", ""), d.get("note", "")]
                          for d in details]
            create_table(self.doc, len(table_data), 5, data=table_data, headers=headers,
                         col_widths=[Cm(2.5), Cm(3), Cm(4.5), Cm(3), Cm(3)])

    def _sub_revenue(self):
        add_subsection_title(self.doc, "2", "총 관객 수 및 수익 결산")
        rev = self.data.get("revenue", {})
        if rev.get("total_visitors"):
            add_sub2_title(self.doc, "1", f"총 관객 수 {rev['total_visitors']}")
            if rev.get("daily_average"):
                add_bullet_main(self.doc, "일평균 관객", rev["daily_average"])

        if rev.get("total_revenue"):
            add_sub2_title(self.doc, "2", f"총 수입 {rev['total_revenue']}")
            if rev.get("ticket_revenue"):
                add_bullet_main(self.doc, "입장 수입", rev["ticket_revenue"])

    def _sub_visitor_composition(self):
        add_subsection_title(self.doc, "3", "관객 구성")
        vc = self.data.get("visitor_composition", {})
        comp = self.data.get("comparison", {})

        # ── 주차별 관객 추이 (비교형: 같은 유형 평균·마지막 기준선) ──
        weekly = vc.get("weekly_visitors", {})
        if weekly and len(weekly) >= 2:
            ref_lines = comp.get("weekly_ref", []) or None
            chart_path = create_weekly_visitors_chart(
                weekly, title="주차별 관객 추이", ref_lines=ref_lines)
            self.temp_files.append(chart_path)
            add_image(self.doc, chart_path, is_chart=True)
            # 서술 (차트→서술 원칙)
            vals = list(weekly.values())
            peak_week = max(weekly, key=weekly.get)
            note = (f"주차별 관객은 {peak_week}에 최고 {max(vals):,}명을 기록함.")
            wref = comp.get("weekly_ref", [])
            if wref:
                ref_txt = ", ".join(f"{lbl} 주당 {v:,.0f}명" for v, lbl, _ in wref)
                note += f" 점선은 비교 기준({ref_txt})임."
            add_paragraph(self.doc, note, size=Fonts.BODY, space_after=Pt(6),
                          line_spacing=1.5, first_line_indent=Cm(0.5))

        # ── 입장권별 구성 + 유료/무료 비율 (자기완결 도넛 2개, 한 줄) ──
        ticket_type = vc.get("ticket_type", {})
        if ticket_type:
            comp_chart = create_media_composition_chart(
                ticket_type, title="입장권별 관객 구성", unit="명")
            # 유료/무료(초대): 무료 = 초대권, 유료 = 나머지
            free = ticket_type.get("초대권", 0) or 0
            paid = sum(v for k, v in ticket_type.items() if k != "초대권")
            paidfree_chart = None
            if paid > 0 or free > 0:
                paidfree_chart = create_media_composition_chart(
                    {"유료": paid, "무료·초대": free}, title="유료·무료 비율", unit="명")
            charts = [c for c in (comp_chart, paidfree_chart) if c]
            for c in charts:
                self.temp_files.append(c)
            if len(charts) == 2:
                add_images_2col(self.doc, charts, img_width=Cm(8))
            elif charts:
                add_image(self.doc, charts[0], is_chart=True)
            # 서술 (차트→서술 원칙)
            top = max(ticket_type, key=ticket_type.get)
            tt_total = sum(ticket_type.values()) or 1
            note = (f"입장권별로는 {top}이(가) {ticket_type[top]:,}명"
                    f"({ticket_type[top]/tt_total*100:.0f}%)으로 가장 큰 비중을 차지함.")
            if paid + free > 0:
                note += f" 유료 관객은 {paid:,}명({paid/(paid+free)*100:.0f}%)임."
            add_paragraph(self.doc, note, size=Fonts.BODY, space_after=Pt(6),
                          line_spacing=1.5, first_line_indent=Cm(0.5))

        for item in vc.get("ticket_analysis", []):
            if item.startswith("→"):
                add_arrow_note(self.doc, item[1:].strip())
            elif item.startswith("-"):
                add_bullet_sub(self.doc, item[1:].strip())
            else:
                add_bullet_main(self.doc, None, item, bold_value=True, underline_value=True)

    # ─── V. 홍보 ───

    def _has_promotion_data(self):
        promo = self.data.get("promotion", {})
        has_promo = any(promo.get(k, "") for k in ["advertising", "press_release", "web_invitation", "newsletter", "sns", "other"])
        press = self.data.get("press_coverage", {})
        has_press = bool(press.get("print_media")) or bool(press.get("online_media"))
        has_membership = bool(self.data.get("membership", ""))
        return has_promo or has_press or has_membership

    def _section_5_promotion(self):
        add_section_title(self.doc, "V", "홍보 방식 및 언론 보도")

        # 홍보 방식
        add_subsection_title(self.doc, "1", "홍보 방식")
        promo = self.data.get("promotion", {})
        num = 1
        for key, label in [("advertising", "광고"), ("press_release", "보도자료"),
                           ("web_invitation", "웹 초청장"), ("newsletter", "뉴스레터"),
                           ("sns", "SNS"), ("other", "그 외")]:
            content = promo.get(key, "")
            if content:
                add_sub2_title(self.doc, num, label)
                for line in content.split("\n"):
                    line = line.strip()
                    if line:
                        add_bullet_main(self.doc, None, line)
                num += 1

        # 언론보도
        add_subsection_title(self.doc, "2", "언론보도 리스트")
        press = self.data.get("press_coverage", {})
        if press.get("print_media"):
            add_sub2_title(self.doc, "1", "일간지 및 월간지")
            headers = ["매체명", "일자", "제목", "비고"]
            table_data = [[i.get("outlet", ""), i.get("date", ""), i.get("title", ""), i.get("note", "")]
                          for i in press["print_media"]]
            create_table(self.doc, len(table_data), 4, data=table_data, headers=headers,
                         col_widths=[Cm(1.3), Cm(1.3), Cm(9), Cm(4.4)])

        if press.get("online_media"):
            add_sub2_title(self.doc, "2", "온라인 매체")
            headers = ["매체명", "일자", "제목", "URL"]
            table_data = [[i.get("outlet", ""), i.get("date", ""), i.get("title", ""), i.get("url", "")]
                          for i in press["online_media"]]
            create_table(self.doc, len(table_data), 4, data=table_data, headers=headers,
                         col_widths=[Cm(1.5), Cm(1.5), Cm(7.5), Cm(5.5)])

        # 멤버십
        if self.data.get("membership"):
            add_subsection_title(self.doc, "3", "멤버십 커뮤니케이션")
            add_paragraph(self.doc, self.data["membership"], size=Fonts.BODY, line_spacing=1.4)

        # v3: 인라인 분석 (홍보)
        self._insert_section_insights("promotion")

    # ─── VI. Executive Summary ───

    def _section_6_evaluation(self):
        add_section_title(self.doc, "VI", "Executive Summary")

        # 1. 핵심 수치 종합표 — 중립적·정밀한 표가 보고서엔 가장 읽기 쉬움
        #    (v5.3.62: 차트 과잉을 정리하고 표로 환원. 롤리팝·재정패널·산점도·
        #     트렌드·유사막대는 보고서에서 제외 — 분석은 워크스페이스 뷰에서 제공)
        self._insert_summary_metrics_table()

        # 2. 종합 의견 (LLM)
        add_subsection_title(self.doc, "2", "종합 의견")
        self._insert_section_insights("evaluation")

        # 3. 관객 반응 종합 (LLM) — 후기가 있을 때만 자동 표시
        self._insert_audience_response()

        # 4. 데이터 도출 평가 항목 (자동 산출, 항목 있을 때만 표시)
        evaluation = self.data.get("evaluation", {})
        reviews = self.data.get("visitor_reviews", [])
        positive_reviews = [r for r in reviews if r.get("category", "").strip() in ("긍정", "긍정적")]
        negative_reviews = [r for r in reviews if r.get("category", "").strip() in ("부정", "부정적", "건의", "불만")]

        positive = evaluation.get("positive", [])
        negative = evaluation.get("negative", [])
        improvements = evaluation.get("improvements", [])

        if positive or negative or improvements or positive_reviews or negative_reviews:
            add_subsection_title(self.doc, "4", "데이터 도출 평가 항목")

            sub_num = 1
            if positive or positive_reviews:
                add_sub2_title(self.doc, sub_num, "긍정 평가")
                for item in positive:
                    add_bullet_main(self.doc, None, item)
                if positive_reviews:
                    add_paragraph(self.doc, "", space_before=Pt(4))
                    headers = ["분류", "상세 내용(인용)", "출처"]
                    table_data = [[r.get("category", "긍정"), r.get("content", ""), r.get("source", "")]
                                  for r in positive_reviews]
                    create_table(self.doc, len(table_data), 3, data=table_data, headers=headers,
                                 col_widths=[Cm(1.25), Cm(11.75), Cm(2)])
                sub_num += 1

            if negative or negative_reviews:
                add_sub2_title(self.doc, sub_num, "부정 평가")
                for item in negative:
                    add_bullet_main(self.doc, None, item)
                if negative_reviews:
                    add_paragraph(self.doc, "", space_before=Pt(4))
                    headers = ["분류", "상세 내용(인용)", "출처"]
                    table_data = [[r.get("category", "부정"), r.get("content", ""), r.get("source", "")]
                                  for r in negative_reviews]
                    create_table(self.doc, len(table_data), 3, data=table_data, headers=headers,
                                 col_widths=[Cm(1.25), Cm(11.75), Cm(2)])
                sub_num += 1

            if improvements:
                add_sub2_title(self.doc, sub_num, "개선 방안")
                for item in improvements:
                    add_bullet_main(self.doc, None, item)

    def _insert_audience_response(self):
        """관객 반응 종합 — LLM이 큐레이터 선별 후기에서 추출한 테마 분석.

        후기가 없거나 LLM이 빈 결과를 반환하면 섹션 자체를 생략.
        통계적 대표성이 없음을 명시 (큐레이터 선별 표본).
        """
        llm_audience = (self.llm_sections.get("audience_response") or "").strip()
        if not llm_audience:
            return

        add_subsection_title(self.doc, "3", "관객 반응 종합")
        for para_text in llm_audience.split("\n\n"):
            para_text = para_text.strip()
            if para_text:
                add_paragraph(self.doc, para_text, size=Fonts.BODY,
                              space_after=Pt(6), line_spacing=1.5,
                              first_line_indent=Cm(0.5))


# ──────────────────────────────────────────────
# 편의 함수
# ──────────────────────────────────────────────

def generate_report(data, output_path):
    generator = ExhibitionReportGenerator(data)
    return generator.generate(output_path)

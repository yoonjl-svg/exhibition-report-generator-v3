"""
일민미술관 전시보고서 문서 스타일 정의
- 기존 구글 문서 기반 보고서 형식을 엄격히 재현
- 제목 체계: I. → 1. → 1) → ① ② ③
- 전시 개요: 불릿 리스트 (● / -)
- 페이지 번호: 우측 하단
"""

from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


# ──────────────────────────────────────────────
# 색상 팔레트
# ──────────────────────────────────────────────
class Colors:
    BLACK = RGBColor(0x00, 0x00, 0x00)
    WHITE = RGBColor(0xFF, 0xFF, 0xFF)
    DARK_GRAY = RGBColor(0x33, 0x33, 0x33)
    MEDIUM_GRAY = RGBColor(0x66, 0x66, 0x66)
    LIGHT_GRAY = RGBColor(0xBF, 0xBF, 0xBF)
    TABLE_HEADER_BG = "D9D9D9"  # hex for XML
    BLUE = RGBColor(0x00, 0x70, 0xC0)  # 화살표 주석용


# ──────────────────────────────────────────────
# 폰트 설정 (Noto Sans CJK KR)
# ──────────────────────────────────────────────
class Fonts:
    KR = "Noto Sans CJK KR"
    EN = "Noto Sans"

    # 크기 (구글 문서 기반 보고서 실측치)
    TOC_TITLE = Pt(16)         # 목차 페이지 제목
    TOC_ITEM = Pt(11)          # 목차 항목
    SECTION_TITLE = Pt(14)     # I. 전시 개요
    SUBSECTION_TITLE = Pt(12)  # 1. 전시
    SUB2_TITLE = Pt(11)        # 1) 1전시실
    DETAIL_TITLE = Pt(10)      # ① 도면
    BODY = Pt(10)              # 본문
    BODY_SMALL = Pt(9)         # 작은 본문
    TABLE_HEADER = Pt(9)       # 표 헤더
    TABLE_CELL = Pt(9)         # 표 셀
    BULLET_MAIN = Pt(10)       # ● 불릿 항목
    BULLET_SUB = Pt(10)        # - 하위 불릿
    CAPTION = Pt(8)            # 이미지 캡션
    PAGE_NUMBER = Pt(9)        # 페이지 번호


# ──────────────────────────────────────────────
# 페이지 설정
# ──────────────────────────────────────────────
class PageSetup:
    WIDTH = Cm(21.0)   # A4
    HEIGHT = Cm(29.7)
    TOP_MARGIN = Cm(2.54)
    BOTTOM_MARGIN = Cm(2.54)
    LEFT_MARGIN = Cm(2.54)
    RIGHT_MARGIN = Cm(2.54)


# ──────────────────────────────────────────────
# 기본 문서 설정
# ──────────────────────────────────────────────

def setup_document(doc):
    """문서 기본 설정"""
    section = doc.sections[0]
    section.page_width = PageSetup.WIDTH
    section.page_height = PageSetup.HEIGHT
    section.top_margin = PageSetup.TOP_MARGIN
    section.bottom_margin = PageSetup.BOTTOM_MARGIN
    section.left_margin = PageSetup.LEFT_MARGIN
    section.right_margin = PageSetup.RIGHT_MARGIN
    return section


def set_run_font(run, size=Fonts.BODY, bold=False, italic=False,
                 color=Colors.BLACK, underline=False):
    """Run에 Noto Sans CJK 폰트 적용"""
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    run.font.underline = underline
    run.font.name = Fonts.EN
    # 한글(eastAsia) 폰트 설정
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{Fonts.KR}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), Fonts.KR)


# ──────────────────────────────────────────────
# 문단 추가 함수들
# ──────────────────────────────────────────────

def add_paragraph(doc, text="", size=Fonts.BODY, bold=False,
                  alignment=WD_ALIGN_PARAGRAPH.LEFT,
                  space_before=Pt(0), space_after=Pt(4),
                  color=Colors.BLACK, line_spacing=1.15,
                  underline=False, first_line_indent=None,
                  left_indent=None):
    """범용 문단 추가"""
    para = doc.add_paragraph()
    para.alignment = alignment
    pf = para.paragraph_format
    pf.space_before = space_before
    pf.space_after = space_after
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    if left_indent is not None:
        pf.left_indent = left_indent

    if text:
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color, underline=underline)
    return para


def add_horizontal_rule(doc, color="AAAAAA", size="3"):
    """수평선(가로줄) 추가
    color: 선 색상 (기본 연한 회색)
    size: 선 두께 (기본 얇은 선, 8분의 1 포인트 단위)
    """
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    # 문단 하단 테두리로 수평선 구현
    pPr = para._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    return para


# ──────────────────────────────────────────────
# 제목 체계: I. → 1. → 1) → ① ② ③
# ──────────────────────────────────────────────

ROMAN_NUMERALS = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII"]
CIRCLED_NUMBERS = ["①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧", "⑨", "⑩"]


def add_section_title(doc, roman_num, title):
    """대제목: I. 전시 개요"""
    para = add_paragraph(
        doc, f"{roman_num}. {title}",
        size=Fonts.SECTION_TITLE, bold=True,
        space_before=Pt(20), space_after=Pt(10),
        line_spacing=1.3
    )
    return para


def add_subsection_title(doc, number, title, suffix=""):
    """소제목: 1. 전시 또는 1. 전시 연계 프로그램 - 총 8개 ...
    suffix는 제목 뒤에 붙는 추가 텍스트 (일반 굵기)
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(14)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.3

    run = para.add_run(f"{number}. {title}")
    set_run_font(run, size=Fonts.SUBSECTION_TITLE, bold=True)

    if suffix:
        run2 = para.add_run(suffix)
        set_run_font(run2, size=Fonts.BODY, bold=False)

    return para


def add_sub2_title(doc, number, title):
    """하위 소제목: 1) 1전시실  또는  2) 프로그램 운영 사진"""
    return add_paragraph(
        doc, f"{number}) {title}",
        size=Fonts.SUB2_TITLE, bold=True,
        space_before=Pt(10), space_after=Pt(4),
        line_spacing=1.3
    )


def add_detail_title(doc, circled_num, title):
    """상세 항목: ① 도면  ② 전경 사진"""
    return add_paragraph(
        doc, f"{circled_num} {title}",
        size=Fonts.DETAIL_TITLE, bold=True,
        space_before=Pt(8), space_after=Pt(4),
        line_spacing=1.3
    )


# ──────────────────────────────────────────────
# 불릿 리스트 (● / -)
# ──────────────────────────────────────────────

def add_bullet_main(doc, label, value, bold_value=False, underline_value=False):
    """메인 불릿: ● 전시 제목: 《하이퍼 옐로우》"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = para.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.4
    pf.left_indent = Cm(0.5)

    run_bullet = para.add_run("● ")
    set_run_font(run_bullet, size=Fonts.BULLET_MAIN, bold=True)

    if label:
        run_label = para.add_run(f"{label}: ")
        set_run_font(run_label, size=Fonts.BULLET_MAIN, bold=True)

    run_value = para.add_run(str(value))
    set_run_font(run_value, size=Fonts.BULLET_MAIN,
                 bold=bold_value, underline=underline_value)
    return para


def add_bullet_sub(doc, text):
    """하위 불릿: - 세부 내용"""
    return add_paragraph(
        doc, f"- {text}",
        size=Fonts.BULLET_SUB,
        space_before=Pt(1), space_after=Pt(1),
        line_spacing=1.4,
        left_indent=Cm(1.2)
    )


def add_arrow_note(doc, text):
    """화살표 주석: → 파란색 텍스트"""
    return add_paragraph(
        doc, f"→ {text}",
        size=Fonts.BODY, bold=False,
        color=Colors.BLUE,
        space_before=Pt(2), space_after=Pt(4),
        left_indent=Cm(1.0),
        line_spacing=1.3
    )


# ──────────────────────────────────────────────
# 표 스타일
# ──────────────────────────────────────────────

def create_table(doc, rows, cols, data=None, headers=None,
                 col_widths=None, header_bg=True):
    """스타일 표 생성 (회색 헤더 + 검정 테두리)"""
    total_rows = rows + (1 if headers else 0)
    table = doc.add_table(rows=total_rows, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # 열 너비
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = width

    # 테두리
    _set_table_borders(table)

    # 헤더 행
    if headers:
        for i, header_text in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(header_text)
            set_run_font(run, size=Fonts.TABLE_HEADER, bold=True)
            if header_bg:
                _set_cell_bg(cell, Colors.TABLE_HEADER_BG)
            _set_cell_vertical_center(cell)

    # 데이터
    if data:
        start = 1 if headers else 0
        for r, row_data in enumerate(data):
            for c, cell_text in enumerate(row_data):
                cell = table.rows[start + r].cells[c]
                cell.text = ""
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.space_before = Pt(2)
                para.paragraph_format.space_after = Pt(2)
                run = para.add_run(str(cell_text))
                set_run_font(run, size=Fonts.TABLE_CELL)
                _set_cell_vertical_center(cell)

    return table


def create_table_left_aligned(doc, rows, cols, data=None, headers=None,
                               col_widths=None, first_col_bold=False):
    """좌측 정렬 표 (첫 열 굵게 옵션)"""
    table = create_table(doc, rows, cols, data=data, headers=headers,
                         col_widths=col_widths)

    # 데이터 셀을 좌측 정렬로 변경
    start = 1 if headers else 0
    if data:
        for r, row_data in enumerate(data):
            for c, _ in enumerate(row_data):
                cell = table.rows[start + r].cells[c]
                for para in cell.paragraphs:
                    if c == 0:
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    else:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    if first_col_bold and c == 0:
                        for run in para.runs:
                            run.font.bold = True

    return table


# ──────────────────────────────────────────────
# 표 내부 헬퍼
# ──────────────────────────────────────────────

def _set_cell_bg(cell, hex_color):
    """셀 배경색"""
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._element.get_or_add_tcPr().append(shading)


def _set_cell_vertical_center(cell):
    """셀 수직 가운데 정렬"""
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)


def _set_table_borders(table, color="000000", size="4"):
    """표 테두리 (검정)"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


# ──────────────────────────────────────────────
# 이미지
# ──────────────────────────────────────────────

# 이미지 크기 상수
class ImageSize:
    SINGLE_MAX_WIDTH = Cm(10)      # 단일 이미지 최대 너비
    SINGLE_MAX_HEIGHT = Cm(15)     # 단일 이미지 최대 높이 (세로형 제한)
    GRID_IMG_WIDTH = Cm(6)         # 2열 그리드 이미지 너비
    CHART_WIDTH = Cm(10)           # 차트 이미지 너비
    POSTER_WIDTH = Cm(8)           # 포스터 이미지 너비


def _get_image_dimensions(image_path):
    """이미지의 원본 가로/세로 비율 반환"""
    try:
        from PIL import Image as PILImage
        with PILImage.open(image_path) as img:
            return img.width, img.height
    except Exception:
        return None, None


def _calc_constrained_size(image_path, max_width, max_height=None):
    """최대 너비/높이 내에서 비율 유지한 크기 계산. width만 반환."""
    w, h = _get_image_dimensions(image_path)
    if w is None or h is None:
        return max_width

    aspect = w / h
    result_width = max_width

    # 세로형 이미지: 높이 제한 적용
    if max_height and h > w:
        width_from_height = Cm(max_height.cm * aspect)
        if width_from_height < max_width:
            result_width = width_from_height

    return result_width


def add_image(doc, image_path, width=None, caption=None, is_chart=False):
    """이미지 추가 (가운데 정렬, 크기 자동 조절)"""
    import os
    if not os.path.exists(image_path):
        return doc.add_paragraph()

    if width is None:
        if is_chart:
            width = ImageSize.CHART_WIDTH
        else:
            width = _calc_constrained_size(
                image_path,
                ImageSize.SINGLE_MAX_WIDTH,
                ImageSize.SINGLE_MAX_HEIGHT
            )

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run()
    run.add_picture(image_path, width=width)

    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap.add_run(caption)
        set_run_font(cap_run, size=Fonts.CAPTION, color=Colors.MEDIUM_GRAY, italic=True)

    return para


def add_images_auto(doc, image_paths):
    """이미지 자동 배치: 1개면 중앙 단독, 2개 이상이면 2열 그리드.
    홀수 개일 경우 마지막 이미지는 중앙에 단독 배치."""
    import os
    valid = [p for p in image_paths if os.path.exists(p)]
    if not valid:
        return None

    # 1개: 중앙 단독
    if len(valid) == 1:
        return add_image(doc, valid[0])

    # 2개 이상: 짝수분을 2열 그리드로, 홀수 마지막은 단독 중앙
    paired = valid if len(valid) % 2 == 0 else valid[:-1]
    last_solo = valid[-1] if len(valid) % 2 != 0 else None

    result = None
    if paired:
        result = _add_images_grid(doc, paired, ImageSize.GRID_IMG_WIDTH)

    if last_solo:
        add_image(doc, last_solo)

    return result


def _add_images_grid(doc, image_paths, img_width):
    """2열 이미지 그리드 (테두리 없는 표, 간격 최소화)"""
    import os
    cols = 2
    rows_needed = (len(image_paths) + cols - 1) // cols
    table = doc.add_table(rows=rows_needed, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    _remove_table_borders(table)

    # 셀 여백을 최소화하여 이미지가 거의 붙도록 설정
    _set_table_cell_margins(table, top=0, bottom=0, start=28, end=28)

    for idx, img_path in enumerate(image_paths):
        if not os.path.exists(img_path):
            continue
        row_idx = idx // cols
        col_idx = idx % cols
        cell = table.rows[row_idx].cells[col_idx]
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        try:
            run = para.add_run()
            run.add_picture(img_path, width=img_width)
        except Exception:
            run = para.add_run("[이미지]")
            set_run_font(run, size=Fonts.CAPTION, color=Colors.LIGHT_GRAY)

    return table


def _set_table_cell_margins(table, top=0, bottom=0, start=0, end=0):
    """표 전체의 셀 여백 설정 (단위: twips, 1cm = 567twips)"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>')
    cell_margin = parse_xml(
        f'<w:tblCellMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{start}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{end}" w:type="dxa"/>'
        f'</w:tblCellMar>'
    )
    tblPr.append(cell_margin)


# 하위 호환: 기존 add_images_2col 호출을 유지
def add_images_2col(doc, image_paths, img_width=Cm(7)):
    """2열 이미지 그리드 — add_images_auto 사용 권장"""
    return add_images_auto(doc, image_paths)


def _remove_table_borders(table):
    """표 테두리 제거"""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(
        f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    # 셀 단위 테두리도 제거
    for row in table.rows:
        for cell in row.cells:
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            tcBdr = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="none" w:sz="0" w:space="0"/>'
                f'  <w:left w:val="none" w:sz="0" w:space="0"/>'
                f'  <w:bottom w:val="none" w:sz="0" w:space="0"/>'
                f'  <w:right w:val="none" w:sz="0" w:space="0"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBdr)


# ──────────────────────────────────────────────
# 페이지 나누기 / 페이지 번호 (우측 하단)
# ──────────────────────────────────────────────

def add_page_break(doc):
    """페이지 나누기"""
    doc.add_page_break()


def add_page_numbers_right(doc):
    """페이지 번호 — 우측 하단"""
    section = doc.sections[0]
    footer = section.footer
    footer.is_linked_to_previous = False
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    run = para.add_run()
    fld1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fld1)

    run2 = para.add_run()
    instr = parse_xml(
        f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instr)

    run3 = para.add_run()
    fld2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fld2)

    set_run_font(run, size=Fonts.PAGE_NUMBER, color=Colors.MEDIUM_GRAY)
